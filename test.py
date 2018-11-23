#!/usr/bin/env python
# coding: utf-8

# # Hack [nairaland.com](https://www.nairaland.com/)

# In[1]:


import re
import os
import sys
import codecs
import logging

from operator import itemgetter
from itertools import filterfalse
from collections import OrderedDict, namedtuple, Counter

import bs4
import docx
import requests
import requests.exceptions as rqe
import openpyxl as OP

from pywebber import Ripper

USERHOME = os.path.abspath(os.path.expanduser('~'))
DESKTOP = os.path.abspath(USERHOME + '/Desktop/')
BASE_DIR = STATUS_DIR = os.path.join(DESKTOP ,"hack-nairaland")

if not os.path.exists(BASE_DIR):
    os.mkdir(BASE_DIR)


# In[2]:


class Error(Exception):
    pass

class NonExistentNairalandUser(Error):
    pass


# In[3]:


def new_logger(log_file_name):
    FORMATTER = logging.Formatter("%(asctime)s:%(funcName)s:%(levelname)s\n%(message)s")
    # console_logger = logging.StreamHandler(sys.stdout)
    file_logger = logging.FileHandler(log_file_name)
    file_logger.setFormatter(FORMATTER)

    logger = logging.getLogger(log_file_name)
    logger.setLevel(logging.DEBUG)
    logger.addHandler(file_logger)
    logger.propagate = False
    return logger

PARSE_BR_TAG_LOGGER = new_logger('log_html_br_tag.log')
PARSE_COMMENT_BLOCK_LOGGER = new_logger('log_parse_comment_block.log')
FORMAT_COMMENTS_LOGGER = new_logger('log_format_comments.log')

def unique_everseen(iterable, key=None):
    """List unique elements, preserving order. Remember all elements ever seen.
    source: https://docs.python.org/3/library/itertools.html#itertools-recipes"""
    seen = set()
    seen_add = seen.add
    if key is None:
        for element in filterfalse(seen.__contains__, iterable):
            seen_add(element)
            yield element
    else:
        for element in iterable:
            k = key(element)
            if k not in seen:
                seen_add(k)
                yield element

def dict_to_string(dictionary):
    """Flatten a dictionary into a single string"""
    if isinstance(dictionary, dict):
        return '\n'.join([' says\n'.join([key, value]) for key, value in dictionary.items()])
    return None

def parse_html_br_tag_content(break_tag):
    """Get content of the next and previous sibling of a <br/> tag

    Parameters
    -----------
    BeautifulSoup
        BeautifulSoup object of <br/> tag

    Returns
    -------
    tuple

    Notes
    ------
    1. For each <br/> tag we run the .string method on its next and previous siblings
    2. If a proper string is encountered, it is returned.
    3. If <br/> is encountered, None is returned.
    4. If any other string which result in an error is encountered, None is returned
    """
    p_sibling = break_tag.previous_sibling
    n_sibling = break_tag.next_sibling

#     PARSE_BR_TAG_LOGGER.debug("previous sibling\n{}".format(p_sibling))
#     PARSE_BR_TAG_LOGGER.debug("next sibling\n{}".format(n_sibling))

    return_value = [None, None]
    try:
        return_value[0] = n_sibling.string.strip().strip("\n:")
    except AttributeError:
        pass
    try:
        return_value[1] = p_sibling.string.strip().strip("\n:")
    except AttributeError:
        pass
    # return ("*{}*, *{}*".format(return_value[0], return_value[1]))
    return tuple(return_value)

def join_tuples(list_of_tuples):
    """Join a list of tuples into a list eliminating None and
    duplicates.

    Parameters
    -----------
    list
        A list of tuples whose elements are strings.

    Returns
    --------
    list
        A list of unique strings in the input list

    Notes
    ------
    unique_everseen removes duplicates.
    This is needed in cases where a next_sibling and a
    previous_sibling point to the same string.
    """
#     remove_nones = [[filter(lambda x: x is not None, each)] for each in list_of_tuples]
    
#     # a more explicit way
#     remove_nones2 = []
#     for each_tuple in list_of_tuples:
#         l = []
#         for each_string in each_tuple:
#             if each_string is not None:
#                 l.append(each_string)
#         remove_nones2.append(l)
#     print(remove_nones2)

    # use list comprehension for speed
    remove_nones = [
        [each_string for each_string in each_tuple if each_string is not None] for each_tuple in list_of_tuples
    ]
    phrase_collection = [phrase.strip().strip("\n:") for each in remove_nones for phrase in each]
    return "\n".join(unique_everseen(phrase_collection))

def format_comments(bs4_comment_block_object):
    """Format a comment block into proper paragraphs

    Parameters
    ------------
    BeautifulSoup
        BeautifulSoup object of comment block

    Returns
    --------
    str
        A properly paragraphed string
    """

#     FORMAT_COMMENTS_LOGGER.debug(bs4_comment_block_object.prettify())

    comment = []
    break_tags = bs4_comment_block_object.find_all('br')

    if break_tags == []:
        return bs4_comment_block_object.text
    for each in break_tags:
        content = parse_html_br_tag_content(each)
        comment.append(content)

    return_string = join_tuples(comment)
    return return_string

def parse_comment_block(bs4_comment_block_object):
    """Return quoted string.
    
    Parameters
    -----------
    BeautifulSoup
        BeautifulSoup object of a quoted string block

    Returns
    --------
    quoted : OrderedDict()
        quoted string content
    bs4_comment_block_object : BeautifulSoup
        Input comment block stripped of all <b> tags
        
    Notes
    ------
    Every comment block must be parsed with this function.
    This function also has a side effect of producing a properly formatted html of all comments it encounters.
    """
    
#     PARSE_COMMENT_BLOCK_LOGGER.debug(bs4_comment_block_object.prettify())

    save_dir = os.path.join(BASE_DIR, "comment_block")
    if os.path.exists(save_dir) is False:
        os.mkdir(save_dir)

    # Side effect
    save = os.path.join(save_dir, "all_page_comments.html")
    with codecs.open(save, 'a+', encoding='utf-8') as f:
        f.write(bs4_comment_block_object.prettify())
        f.write("<div class='dropdown-divider'></div>")

    collected_quotes = OrderedDict()
    return_val = namedtuple('ParsedComment', ['focus_user_comment', 'quotes_ordered_dict'])
    blockquotes = bs4_comment_block_object.find_all('blockquote')

    if blockquotes == []:
        pass
    else:
        for blockquote in blockquotes:
            try:
                commenter = blockquote.find('b').text
            except AttributeError:
                commenter = 'Anonymous'

            collected_quotes[commenter] = format_comments(blockquote).strip().strip("\n:")
            blockquote.decompose() # remove the block from the tree

    return_val.focus_user_comment = format_comments(bs4_comment_block_object).strip().strip("\n:")
    return_val.quotes_ordered_dict = collected_quotes
    return return_val

def sort_dictionary_by_value(dictionary_to_sort):
    """
    Return list of dictionary keys where the items are sorted on the values in descending order.
    
    e.g sort_dictionary_by_value({5:'goat', 10:'cat', 1:'dog'}) returns [5, 1, 10] since the values
    sort to ['goat', 'dog', 'cat']
    """
    if dictionary_to_sort is None:
        return
    ordered_dictionary = sorted(dictionary_to_sort.items(), key=itemgetter(1), reverse=True)
    sorted_dictionary_list = [i[0] for i in ordered_dictionary]
    return sorted_dictionary_list


# In[4]:


class PostCollector:
    """
    Scrap a nairaland post

    Parameters
    ----------
    str
        Post url
    """

    def __str__(self):
        return "PostCollector: {}".format(self.base_url)

    def __init__(self, base_url, refresh=True):
        self.save_path = os.path.join(BASE_DIR, 'page_rips_post')
        self.base_url = base_url # Page (0) of the post
        self.refresh = refresh
        self.title = self.base_url.split('/')[-1]
        if not os.path.exists(self.save_path):
            os.mkdir(self.save_path)

    def max_page(self):
        """Returns the maximum number of pages of comments, starting from a zero index."""
        stop = 0
        while True:
            print("Whiling in a loop")
            if self._check_if_url_exists_and_is_valid("{}/{}".format(self.base_url, stop)): # check if next url exists
                stop += 1
            else:
                break
        return stop

    @staticmethod
    def _check_if_url_exists_and_is_valid(url): # ConnectionError happens here
        r = requests.head(url)
        return r.status_code == 200

    def _scrap_comment_for_single_page(self, page_url):
        """Return comments and commenters on a single post page

        Returns
        --------
        OrderedDict
            Dictionary of {commenter : comments}
        """

        # User posts are contained in a table with summary='posts' attribute.
        # Each commenter name is contained inside a <tr>
        # Each comment is contained in <tr> just below the name of the commenter
        soup = Ripper(page_url, parser='html5lib', save_path=self.save_path, refresh=self.refresh).soup

        # Handle supposed anomaly by decomposing all such occurrences from the tree
        for each in soup.find_all('td', class_="l pu pd"):
            each.parent.decompose()
        rows = soup.find('table', summary='posts').find_all('tr')

        return_val = OrderedDict()
        for i in range(0, len(rows), 2):

            topic_classes = ['bold l pu', 'bold l pu nocopy'] # topic div should be either of these classes
            for class_ in topic_classes:
                try:
                    moniker = rows[i].find('td', class_=class_).find('a', href=True, class_=True).text.strip()
                    break
                except AttributeError:
                    pass
                moniker = "Nobody" # set to nobody after exhausting all options. We cannot use finally in this case

            comment_classes = ['l w pd', 'l w pd nocopy'] # comment div should be either of these classes            
            for class_ in comment_classes:
                try:
                    comment_block = rows[i+1].find('td', id=True, class_=class_).find('div', class_='narrow')
                    break
                except AttributeError:
                    pass
            parsed_block = parse_comment_block(comment_block)

            # If a moniker already exists (i.e. a user has already commented), append an integer to the
            # present one to differentiate both.
            if moniker not in return_val:
                return_val[moniker] = parsed_block
            else:
                moniker = "{}**{}".format(moniker, i)
                return_val[moniker] = parsed_block
        return return_val

    def scrap_comments_for_range_of_pages(self, start=0, stop=1, __all=False):
        """Get contents for a range of pages from start to stop"""
        if __all == True: # since we're starting from a zero index, we have to subtract 1 from self.max_page()
            stop = self.max_page() - 1
        while start <= stop:
            next_url = "{}/{}".format(self.base_url, start)
            next_page = self._scrap_comment_for_single_page(next_url)
            yield next_page
            start += 1

    def all_commenters(self):
        """Return list of all commenters on a post"""
        # Remember we user ** to separate a moniker and the number of times it is appearing on a post
        return sorted([key.split("**")[0] for each in list(self.scrap_comments_for_range_of_pages(stop=self.max_page())) for key, value in each.items()])

    def unique_commenters(self):
        """Return list of unique commenters on a post"""
        return sorted(set(self.all_commenters()))

    def commenters_activity_summary(self):
        """Return count of number of times a user commented on a post"""
        x = Counter(self.all_commenters())
        print(x)
        print("Sorted dict")
        print(sort_dictionary_by_value(x))
        return 

class UserCommentHistory:
    """
    Grab a user's comment history

    Parameters
    ------------
    user_name : str
        User's name to crawl. Default is 'seun'
    """

    def __str__(self):
        return "UserCommentHistory: {}".format(self.user_post_page)

    def __init__(self, nairaland_moniker, refresh=True):
        self.refresh = refresh
        BASE_URL = 'https://www.nairaland.com'
        self.save_path = os.path.join(BASE_DIR, 'page_rips_user')
        if not os.path.exists(self.save_path):
            os.mkdir(self.save_path)
            
        p = '{}/{}'.format(BASE_URL, nairaland_moniker.lower())
        if self._check_if_url_exists_and_is_valid(p):
            self.user_profile_page = p
            self.user_post_page = '{}/{}/posts'.format(BASE_URL, nairaland_moniker.lower())
        else:
            raise NonExistentNairalandUser("This user does not exist on nairaland.")

    @staticmethod
    def _check_if_url_exists_and_is_valid(url):
        r = requests.head(url)
        return r.status_code == 200
    
    def user_profile(self):
        """Returns a dictionary of the user's profile"""
        pass

    def max_pages(self):
        """Return number of pages of comment for user"""
        soup = Ripper(self.user_post_page, save_path=self.save_path, refresh=True).soup
        pattern = r"\<b\>\s*(\d+)\s*\<\/b\>" # pattern to search for number of pages of comments
        try:
            return int(re.search(pattern, str(soup)).group(1))
        except AttributeError:
            print("Could not find max page")
            pass

    def _scrap_comment_for_single_page(self, page_url):
        """Return comments and commenters on a single post page

        Returns
        --------
        OrderedDict
            Dictionary of {section : namedtuple}
        """

        # User comments are contained in a table with neither summary nor id attribute.
        # Then follows the rows containing the section, topic, and username and the comment itself just below it
        soup = Ripper(page_url, parser='html5lib', save_path=self.save_path, refresh=self.refresh).soup

        for each in soup.find_all('td', class_="l pu pd"):
            each.parent.decompose() # remove these trees as they are unneeded
        rows = soup.find('table', id=False, summary=False).find_all('tr')

        return_val = OrderedDict()
        for i in range(0, len(rows), 2): # go to every second row
            
            topic_classes = ['bold l pu', 'bold l pu nocopy']
            for class_ in topic_classes:
                try:
                    section_topic = rows[i].find('td', class_=class_).find_all('a', href=True, class_=False)
                except AttributeError:
                    pass

            comment_classes = ['l w pd', 'l w pd nocopy']
            for class_ in comment_classes:
                try:
                    comment_block = rows[i+1].find('td', class_=class_).find('div', class_='narrow')
                except AttributeError:
                    pass

            section = section_topic[0].text.strip()
            topic = section_topic[1].text.lstrip("Re:").strip()

            parsed_block = parse_comment_block(comment_block)

            Comm = namedtuple('Comment', ['topic', 'parsed_comment'])
            Comm.topic = topic
            Comm.parsed_comment = parsed_block

            if section not in return_val:
                return_val[section] = Comm
            else:
                section = "{}**{}".format(section, i)
                return_val[section] = Comm
        return return_val

    def scrap_comments_for_range_of_pages(self, start=0, stop=0, _maximum_pages=False):
        """Get contents for a range of pages from start to stop """
        if _maximum_pages:
            stop = self.max_pages() - 1
        while start <= stop:
            next_url = "{}/{}".format(self.user_post_page, start)
            next_page = self._scrap_comment_for_single_page(next_url)
            yield next_page
            start += 1

class TopicCollector:
    """
    Collect topics from a section of nairaland

    Parameters
    -----------
    section : str
        Default section is politics

    Notes
    ------
    The methods in this class are only applicable to section urls
    """

    def __str__(self):
        return "TopicCollector: {}".format(self.base_url)

    def __init__(self, section='politics'):
        self.save_path = os.path.join(BASE_DIR, 'page_rips_section')
        self.section = section
        self.base_url = 'http://www.nairaland.com/{}'.format(self.section)
        if os.path.exists(self.save_path) is False:
            os.mkdir(self.save_path)

    def max_pages(self):
        """Return number of pages in this section

        Returns
        --------
        int
            The number of pages in this section
        """
        soup = Ripper(self.base_url, parser='html5lib', save_path=self.save_path, refresh=True).soup
        number = re.search(r"\(of\s*(\d+)\s*pages\)", soup.text).group(1)
        return int(number)

    def _scrap_topics_for_a_single_page(self, page_url, refresh=True):
        """
        Yield all topics on a page

        Yields
        -------
        namedtuple
            collection of 'poster', 'title', 'url', 'number of comments'
        """
        soup = Ripper(page_url, parser='html5lib', save_path=self.save_path, refresh=True).soup
        post_table = soup.find('table', id=False, summary=False)

        for td in post_table.find_all('td', id=True):
            Post = namedtuple('Post', ['poster', 'title', 'url', 'comments'])

            title_component = td.find('b').find('a', href=True)
            Post.title = title_component.text.strip()
            Post.url = 'http://www.nairaland.com' + title_component.get('href').strip()

            # there is a maximum of 7 <b> tags
            meta_component = td.find('span', class_='s').find_all('b')

            Post.poster = meta_component[0].text.strip()
            Post.comments = int(meta_component[1].text) # count includes the post itself
            yield Post

    def scrap_topics_for_range_of_pages(self, start=0, end=0, _maximum_pages=False):
        """Yield all topics between 'start' and 'end' for a section

        Parameters
        -----------
        int
            Start and end values of section

        Yields
        -------
        tuple
            same yields as for titles()
        """
        if _maximum_pages:
            end = self.max_pages() - 1
        while start <= end:
            next_url = '{}/{}'.format(self.base_url, start)
            yield self._scrap_topics_for_a_single_page(next_url)
            start += 1


# In[ ]:


def export_user_comments_to_html(username=None, max_page=5):
    """Export all of a user's comments data to a html file

    Parameters
    -----------
    str
        Username
    int
        Maximum page count for user's comments (Default is 5 pages of comments)
        loop breaks if we exceed actual count
    """
    
    if not username:
        print("No username provided. Ending")
        return
    else:
        print("Now hacking nairaland. Please wait a few minutes.")
        
    destination_file = os.path.join(BASE_DIR, "comments_{}_{}_pages.html".format(username.lower(), max_page))
    if os.path.exists(destination_file):
        os.remove(destination_file)
    with open(destination_file, 'a+', encoding='utf-8') as f:
        
        # html scaffold
        f.write("<html xmlns='http://www.w3.org/1999/xhtml'>\n")
        f.write("<head>\n")
        f.write("<link rel='stylesheet' href='https://stackpath.bootstrapcdn.com/bootswatch/4.1.3/superhero/bootstrap.min.css'>\n")
        f.write("<meta name='viewport' content='width=device-width, initial-scale=1, shrink-to-fit=no'>\n")
        f.write("<script src='https://stackpath.bootstrapcdn.com/bootstrap/4.1.1/js/bootstrap.min.js' integrity='sha384-smHYKdLADwkXOn1EmN1qk/HfnUcbVRZyYmZ4qpPea6sjB/pTJ0euyQp0Mk8ck+5T' crossorigin='anonymous' async></script>\n")
        f.write("<title>Hack Nairaland - comment history for {}</title>\n".format(username.lower()))
        f.write("</head>\n")
        f.write("<body style='padding-top:5.5rem;'>\n")
        # navbar
        f.write("<nav class='navbar navbar-expand-lg navbar-dark bg-primary fixed-top'>\n")
        f.write("<a class='navbar-brand'>Hack Nairaland</a>\n")
        f.write("<button type='button' class='navbar-toggler my-toggler' data-toggle='collapse' data-target='.navcontent'>\n")
        f.write("<span class='sr-only'>Toggle navigation</span>\n")
        f.write("<span class='navbar-toggler-icon'></span>\n")
        f.write("</button>\n")
        f.write("<div class='collapse navbar-collapse navcontent'>\n")
        f.write("<ul class='nav navbar-nav lefthand-navigation'>\n")
        f.write("<li class='nav-item'><a class='nav-link' href='#' title='Home'>Home</a></li>\n")
        f.write("</ul>\n")
        f.write("</div>\n")
        f.write("</nav>\n")
        # end navbar
        f.write("<div class='container'>\n")
        f.write("<h1>Nairaland comment history for <a href='https://nairaland.com/{0}/posts' target='_blank'>{0}</a></h1>\n".format(username))
        # breadcrumb
        f.write("<nav aria-label='breadcrumb'>\n")
        f.write("<ol class='breadcrumb'>\n")
        f.write("<li class='breadcrumb-item'><a href='#'>Home</a></li>\n")
        f.write("<li class='breadcrumb-item'>The first {} pages</li>\n".format(max_page))
        f.write("</ol>\n")
        f.write("</nav>\n")
        # end breadcrumb

        user = UserCommentHistory(username)
        for page in list(user.scrap_comments_for_range_of_pages(stop=max_page)):
            for section, topic_plus_comment in page.items():
                f.write("<h3>{}</h3>\n".format(section.split('**')[0])) # remove the ** separating section and index
                parsed_comment = topic_plus_comment.parsed_comment
                f.write("<p class='text-success'>{}</p>\n".format(parsed_comment.focus_user_comment))
                quotes = parsed_comment.quotes_ordered_dict
                
                for username, comment in quotes.items():
                    f.write("<h4>{}</h4>\n".format(username))
                    f.write("<p class='text-primary'><em>{}</em></p>\n".format(comment))
                f.write("<div class='dropdown-divider'></div>\n")
            f.write("<hr>\n")
        f.write("</div>\n")
        f.write("</body>")    
    print("Done hacking")


# In[16]:


# To do
def export_user_comments_to_excel(username=None, max_page=5):
    """Export all of a user's comments data to excel

    Parameters
    -----------
    str
        Username
    int
        Maximum page count for user's comments (Default is ..)
        loop breaks if we exceed actual count
    """
    if not username:
        print("No username provided. Ending")
        return
    else:
        print("Now hacking nairaland. Please wait a few minutes.")
        
    user_data = UserCommentHistory(username).scrap_comments_for_range_of_pages(stop=max_page)
    user_comments = list(user_data)
    work_book = OP.Workbook()
    active_sheet = work_book.active
    active_sheet.title = username
    active_sheet['A1'] = "SECTION"
    active_sheet['B1'] = "TOPIC"
    active_sheet['C1'] = "QUOTED"
    active_sheet['D1'] = "USER COMMENT"

    user = UserCommentHistory(username)
        for page in list(user.scrap_comments_for_range_of_pages(stop=max_page)):
            for section, topic_plus_comment in page.items():
                f.write("<h3>{}</h3>\n".format(section.split('**')[0])) # remove the ** separating section and index
                parsed_comment = topic_plus_comment.parsed_comment
                f.write("<p class='text-success'>{}</p>\n".format(parsed_comment.focus_user_comment))
                quotes = parsed_comment.quotes_ordered_dict
                
                for username, comment in quotes.items():
                    f.write("<h4>{}</h4>\n".format(username))
                    f.write("<p class='text-primary'><em>{}</em></p>\n".format(comment))
                f.write("<div class='dropdown-divider'></div>\n")
            f.write("<hr>\n")
        f.write("</div>\n")
        f.write("</body>")    
    print("Done hacking")

    row_number = 2
    for each_comment in user_comments:

        poster = each_comment[0]
        section = each_comment[1]
        topic = each_comment[2]
        quoted = single_string_from_dictionary(each_comment[3])
        comment = each_comment[4]

        active_sheet.cell(row=row_number, column=1, value=poster)
        active_sheet.cell(row=row_number, column=2, value=section)
        active_sheet.cell(row=row_number, column=3, value=topic)
        active_sheet.cell(row=row_number, column=4, value=quoted)
        active_sheet.cell(row=row_number, column=5, value=comment)

        row_number += 1
    work_book.save(os.path.join(BASE_DIR, "comments_{}_{}_pages.xlsx".format(username.lower(), max_page)))

def title_word_count(section='romance', max_page=10):
    """Group words that occur in titles of a section

    Parameters
    -----------
    str
        Section
    int
        Maximum number of section pages to scrap

    Returns
    --------
    word_count : list
        Dictionary of each word and its count
    list
        A list of all words sorted in descending order of frequency
    """
    word_list = []
    section_object = TopicCollector(section=section).titles_links_metadata(end=max_page)
    section_data = list(section_object)
    split_string = r"[\; \, \n \.+\- \( \) - \/ : \? \[ \] \ — –]"

    exclude = ['to', 'a', 'as', 'the', 'you', 'in', 'is', 'i', 'with', 'of', 'an', 'and', 'my',
               'your', 'for', 'on', 'what', 'her', 'this', 'that', 'these', 'those', 'me']

    for each in section_data:
        title = each[3]
        words_in_title = re.split(split_string, title)
        word_list.extend(words_in_title)

    word_list = [x.lower() for x in word_list]
    word_list = filter(lambda x: x not in exclude, word_list)
    word_count = Counter(list(word_list))
    return word_count, sort_dictionary_by_value(word_count)

def export_topics_to_excel(section='romance', start_page=0, end_page=3):
    """Writes all topics between start and end of a section to excel.
    Same output as titles_links_metadata() but written to a excel file
    """
    work_book = OP.Workbook()
    active_sheet = work_book.active
    active_sheet.title = section
    active_sheet['A1'] = "POSTER"
    active_sheet['B1'] = "MONTH"
    active_sheet['C1'] = "YEAR"
    active_sheet['D1'] = "TITLE"
    active_sheet['E1'] = "WEB URL"

    section_object = TopicCollector(section=section)
    section_titles = section_object.titles_links_metadata(start=start_page, end=end_page)
    section_data = list(section_titles)

    row_number = 2

    for each in section_data:
        title = each[0]
        link = each[1]
        month = each[2]
        year = each[3]
        poster = each[4]

        active_sheet.cell(row=row_number, column=1, value=title)
        active_sheet.cell(row=row_number, column=2, value=link)
        active_sheet.cell(row=row_number, column=3, value=month)
        active_sheet.cell(row=row_number, column=4, value=year)
        active_sheet.cell(row=row_number, column=5, value=poster)
        row_number += 1

    fname = "{}_topics_{}_to_{}.xlsx".format(section, start_page, end_page)
    work_book.save(os.path.join(BASE_DIR, fname))

def export_post_to_docx_format(post_url):
    """Export post to word"""

    post = PostCollector(post_url)
    comments_and_commenters = post.view_comments_commenters()

    document = docx.Document()
    document.add_paragraph(post_url)

    for each in comments_and_commenters:
        commenter = each[0]
        comment_block = each[1]

        quotes = single_string_from_dictionary(comment_block[0])
        user_comments = comment_block[1]

        document.add_paragraph().add_run(commenter).bold = True
        document.add_paragraph().add_run('[{}]'.format(quotes)).italic = True
        document.add_paragraph(user_comments)
        document.add_paragraph('*'*50)

    fname = "{}.docx".format(post.post_title)
    document.save(os.path.join(BASE_DIR, fname))


# export_user_comments_to_html(username="seun", max_page=10)

# ## Change the cell type to code to run this cell and view the output
# 
# import textwrap
# 
# user_comments = UserCommentHistory("preccy69")
# for page in list(user_comments.scrap_comments_for_range_of_pages(start=0, stop=1)):
#     for section, topic_plus_comment in page.items():
#         print("\n\n", "*"*40, section, "*"*40)
#         print(topic_plus_comment.topic.upper()) # for differentiation only
# 
#         parsed_comment = topic_plus_comment.parsed_comment # a namedtuple instance
#         print(parsed_comment.focus_user_comment)
# 
#         quotes = parsed_comment.quotes_ordered_dict
#         for username, comment in quotes.items():
#             print(" "*8)
#             print(textwrap.indent(username, "    "))
#             print(textwrap.indent(comment, "    "))
#             
#         print("_"*100)
#     print("\n\n")

# ## Change the cell type to code to run this cell and view the output
# 
# import textwrap
# p = TopicCollector(section='politics')
# for page in p.scrap_topics_for_range_of_pages(end=1):
#     for topic in list(page):
#         print(topic.poster)
#         print(textwrap.indent(topic.title, "    "))
#         print(textwrap.indent(topic.url, "    "))
#         print(textwrap.indent(str(topic.comments), "    "), " comments")
#         print()

# ## Tests

# In[28]:


import os
import sys
import unittest
import types
from unittest import mock

from bs4 import BeautifulSoup

from collections import OrderedDict

from pager import PostCollector, UserCommentHistory, TopicCollector
from utils import parse_br_tag, parse_comment_block_alt

# from . import pager

class TestParseCommentBlock(unittest.TestCase):
    def test_comment_block_parsing(self):
        with open('data_test/test_comment_parser_input.html', 'r+') as rh:
            soup = BeautifulSoup(rh.read(), 'html5lib').find('div', class_='narrow')
        with open('data_test/test_comment_parser_output.txt', 'r+') as r:
            excpected = r.read()
        parsed_data = parse_comment_block_alt(soup)

        print(excpected)
        print()
        print(parsed_data)
        self.assertEqual(parsed_data, excpected)

class TestParsebrTag(unittest.TestCase):
    def test_tag_only_previous_sibling(self):
        tag_block = """
        This tag has previous sibling text BUT no next sibling.
        <br/>
        """
        tag = BeautifulSoup(tag_block, 'html5lib').find('br')
        self.assertEqual(parse_br_tag(tag), "This tag has previous sibling text BUT no next sibling.")

    def test_tag_with_only_next_sibling(self):
        tag_block = """<br/>
        This tag has next sibling text BUT no previous sibling.
        """
        tag = BeautifulSoup(tag_block, 'html5lib').find('br')
        self.assertEqual(parse_br_tag(tag), "This tag has next sibling text BUT no previous sibling.")

    def tag_with_no_previous_or_next_sibling(self):
        tag_block = """
        <br/>
        <br/>
        """
        tag_block2 = """
        <br/>
        <br/>
        <br/>
        """
        tag = BeautifulSoup(tag_block, 'html5lib').find('br')
        self.assertEqual(parse_br_tag(tag), "")
        tag2 = BeautifulSoup(tag_block2, 'html5lib').find('br')
        self.assertEqual(parse_br_tag(tag2), "")

    def test_tag_with_both_next_and_previous_siblings(self):
        tag_block = """
        This tag has previous sibling text
        <br/>
        It also has a next sibling text
        """
        tag = BeautifulSoup(tag_block, 'html5lib').find('br')
        self.assertEqual(parse_br_tag(tag), "This tag has previous sibling text\nIt also has a next sibling text")

class TestPostCollector(unittest.TestCase):
    def setUp(self):
        self.base_url = "some/base/url"

    @mock.patch('pager.os')
    def test_init(self, mock_os):
        # Save directory exists; 'mkdir' NOT called
        mock_os.path.join.return_value = 'joined-path'
        mock_os.path.exists.return_value = True
        obj = PostCollector(self.base_url)
        self.assertEqual(obj.title, 'url')
        self.assertFalse(mock_os.mkdir.called, "Directory creation method called")

        # Save directory doesn't exists; 'mkdir' called
        mock_os.path.join.return_value = 'joined-path'
        mock_os.path.exists.return_value  = False
        PostCollector(self.base_url)
        mock_os.mkdir.assert_called_with('joined-path')

    @mock.patch('pager.requests')
    def test_url_checker(self, mocked_requests):
        check = PostCollector(self.base_url)._check_url_page_exists('whatever-url')

        # Assert that requests head method was called during function execution
        mocked_requests.head.assert_called_with("whatever-url")
        self.assertIsInstance(check, bool)

    @mock.patch('pager.requests')
    def test_max_page(self, mocked_requests):
        obj = PostCollector(self.base_url)
        obj._check_url_page_exists = mock.MagicMock(return_value=False)
        self.assertEqual(obj.max_page(), -1)

    @mock.patch('pager.Ripper')
    def test_get_single_page_content(self, mocked_ripper):
        """Test return object type"""
        obj = PostCollector(self.base_url)
        single_page = obj._get_single_page_content("page-url")
        mocked_ripper.assert_called()
        self.assertIsInstance(single_page, OrderedDict)

    @mock.patch('pager.Ripper')
    def test_scrap_comment_for_page_range(self, mocked_ripper):
        """Test length and type of object returned"""
        obj = PostCollector(self.base_url)
        page_range = obj.scrap_comment_for_page_range(0, 3, False)
        self.assertIsInstance(page_range, types.GeneratorType)
        self.assertEqual(len(list(page_range)), 4)

        # Set maximum page to 3
        obj.max_page = mock.MagicMock(return_value=3)
        page_range = obj.scrap_comment_for_page_range(0, 0, True)
        self.assertEqual(len(list(page_range)), 4)

if __name__ == '__main__':
    unittest.main()


# In[ ]:




