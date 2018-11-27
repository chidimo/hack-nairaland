import re
import os
import sys
import codecs
import logging

from pathlib import Path
from operator import itemgetter
from itertools import filterfalse
from collections import OrderedDict, namedtuple, Counter

import bs4
from docx import Document
import requests
import requests.exceptions as rqe
import openpyxl as OP

from openpyxl.styles import Alignment

from pywebber import Ripper

logging.disable(logging.CRITICAL)

BASE_DIR = Path().resolve()
OUTPUT_DIR = Path.joinpath(BASE_DIR, 'output')
if not Path.exists(OUTPUT_DIR):
    Path.mkdir(OUTPUT_DIR)
TEST_DIR = Path.joinpath(BASE_DIR, 'test-dir')

class Error(Exception):
    pass

class NonExistentNairalandUser(Error):
    pass

class MaximumPageNotFound(Error):
    pass

def new_logger(log_file_name):
    FORMATTER = logging.Formatter("%(asctime)s:%(funcName)s:%(levelname)s\n%(message)s")
    # console_logger = logging.StreamHandler(sys.stdout)
    file_logger = logging.FileHandler(log_file_name)
    file_logger.setFormatter(FORMATTER)

    logger = logging.getLogger(log_file_name)
    logger.setLevel(logging.DEBUG)
    logger.addHandler(file_logger)
    logger.propagate = False
    return logger

PARSE_BR_element_LOGGER = new_logger('log_html_br_element.log')
PARSE_COMMENT_BLOCK_LOGGER = new_logger('log_parse_comment_block.log')
FORMAT_COMMENTS_LOGGER = new_logger('log_format_comments.log')

def check_if_url_exists_and_is_valid(url):
    r = requests.head(url)
    return r.status_code == 200

def unique_everseen(iterable, key=None):
    """List unique elements, preserving order. Remember all elements ever seen.
    source: https://docs.python.org/3/library/itertools.html#itertools-recipes"""
    seen = set()
    seen_add = seen.add
    if key is None:
        for element in filterfalse(seen.__contains__, iterable):
            seen_add(element)
            yield element
    else:
        for element in iterable:
            k = key(element)
            if k not in seen:
                seen_add(k)
                yield element

def get_left_right_of_html_br_element(br_element):
    """Get content of the next and previous sibling of a <br/> tag

    Parameters
    -----------
    BeautifulSoup
        BeautifulSoup object of <br/> tag

    Returns
    -------
    tuple
        A tuple of (previous_sibling_text, next_sibling_text)

    Notes
    ------
    1. For each <br/> tag we run the .string method on its next and previous siblings
    2. If a proper string is encountered, it is returned.
    3. If <br/> is encountered, None is returned.
    4. If any other string which result in an error is encountered, None is returned
    """
    p_sibling = br_element.previous_sibling
    n_sibling = br_element.next_sibling

    PARSE_BR_element_LOGGER.debug("previous sibling\n{}".format(p_sibling))
    PARSE_BR_element_LOGGER.debug("next sibling\n{}".format(n_sibling))

    return_value = [None, None]
    try:
        return_value[1] = n_sibling.string.strip().strip("\n:")
    except AttributeError:
        pass
    try:
        return_value[0] = p_sibling.string.strip().strip("\n:")
    except AttributeError:
        pass
    # return ("*{}*, *{}*".format(return_value[0], return_value[1]))
    return tuple(return_value)

def join_br_tuples(list_of_tuples):
    """Join a list of tuples into a list eliminating None and
    duplicates.

    Parameters
    -----------
    list
        A list of tuples whose elements are strings.

    Returns
    --------
    list
        A list of unique strings in the input list

    Notes
    ------
    unique_everseen removes duplicates.
    This is needed in cases where a next_sibling and a
    previous_sibling point to the same string.
    """
#     remove_nones = [[filter(lambda x: x is not None, each)] for each in list_of_tuples]

#     # a more explicit way
#     remove_nones2 = []
#     for each_tuple in list_of_tuples:
#         l = []
#         for each_string in each_tuple:
#             if each_string is not None:
#                 l.append(each_string)
#         remove_nones2.append(l)
#     print(remove_nones2)

    # use list comprehension for speed
    remove_nones = [
        [each_string for each_string in each_tuple if each_string is not None] for each_tuple in list_of_tuples
    ]
    phrase_collection = [phrase.strip().strip("\n:") for each in remove_nones for phrase in each]
    return "\n".join(unique_everseen(phrase_collection))

def format_comments(bs4_comment_block_object):
    """Format a comment block into proper paragraphs

    Parameters
    ------------
    BeautifulSoup
        BeautifulSoup object of comment block

    Returns
    --------
    str
        A properly paragraphed string
    """

    FORMAT_COMMENTS_LOGGER.debug(bs4_comment_block_object.prettify())

    comment = []
    br_elements = bs4_comment_block_object.find_all('br')

    if br_elements == []:
        return bs4_comment_block_object.text
    for el in br_elements:
        content = get_left_right_of_html_br_element(el) # returns a tuple
        comment.append(content)

    return_string = join_br_tuples(comment)
    return return_string

def parse_comment_block(bs4_comment_block_object):
    """Return quoted string.

    Parameters
    -----------
    BeautifulSoup
        BeautifulSoup object of a quoted string block

    Returns
    --------
    quoted : OrderedDict()
        quoted string content
    bs4_comment_block_object : BeautifulSoup
        Input comment block stripped of all <b> tags

    Notes
    ------
    Every comment block must be parsed with this function.
    This function also has a side effect of producing a properly formatted html of all comments it encounters.
    """

    PARSE_COMMENT_BLOCK_LOGGER.debug(bs4_comment_block_object.prettify())

    save_dir = os.path.join(BASE_DIR, "comment-blocks")
    if os.path.exists(save_dir) is False:
        os.mkdir(save_dir)

    # Side effect
    save = os.path.join(save_dir, "comment-block-collection.html")
    with codecs.open(save, 'a+', encoding='utf-8') as f:
        f.write(bs4_comment_block_object.prettify())
        f.write("End of file.\n\n")

    collected_quotes = OrderedDict()
    output_named_tuple = namedtuple('ParsedComment', ['focus_user_comment', 'quotes_ordered_dict'])
    blockquotes = bs4_comment_block_object.find_all('blockquote')

    # collect comments from other users which were quoted by the focus user
    if blockquotes == []:
        pass
    else:
        for blockquote in blockquotes:
            try:
                commenter = blockquote.find('b').text
            except AttributeError:
                commenter = 'Anonymous'

            collected_quotes[commenter] = format_comments(blockquote).strip().strip("\n:")
            blockquote.decompose() # remove the block from the tree

    # after decomposing all the <blockquote> elements, whatever remains belong to the focus user
    output_named_tuple.focus_user_comment = format_comments(bs4_comment_block_object).strip().strip("\n:")
    output_named_tuple.quotes_ordered_dict = collected_quotes
    return output_named_tuple

def sort_dictionary_by_value(dictionary_to_sort):
    """
    Return list of dictionary keys where the items are sorted on the values in descending order.

    e.g sort_dictionary_by_value({5:'goat', 10:'cat', 1:'dog'}) returns [5, 1, 10] since the values
    sort to ['goat', 'dog', 'cat']
    """
    if dictionary_to_sort is None:
        return
    ordered_dictionary = sorted(dictionary_to_sort.items(), key=itemgetter(1), reverse=True)
    sorted_dictionary_list = [i[0] for i in ordered_dictionary]
    return sorted_dictionary_list

assert sort_dictionary_by_value({5:'goat', 10:'cat', 1:'dog'}) == [5, 1, 10]

class Nairaland(object):
    """The base nairaland class
    sections are contained in a table with class='boards'
    """
    def __init__(self):
        self.site_url = "https://www.nairaland.com/"

        soup = Ripper(self.site_url, parser='html5lib', refresh=True).soup
        boards = soup.find("table", class_="boards")
        links = boards.find_all('a')
        self.sections = {link.text : link.get('href') for each in links}

    def __str__(self):
        return "Nairaland base class"

class PostCollector(Nairaland):
    """
    Scrap a nairaland post

    Parameters
    ----------
    str
        Post url
    """

    def __str__(self):
        return "PostCollector: {}".format(self.post_url)

    def __init__(self, post_url, refresh=True):
        super().__init__()
        self.save_path = os.path.join(OUTPUT_DIR, 'page_rips_post')
        self.post_url = post_url # Page (0) of the post
        self.refresh = refresh
        self.title = self.post_url.split('/')[-1]

        if not os.path.exists(self.save_path):
            os.mkdir(self.save_path)

    def max_page(self):
        """Returns the maximum number of pages of comments, starting from a zero index."""
        stop = 0
        while True:
            if self._check_if_url_exists_and_is_valid("{}/{}".format(self.post_url, stop)): # check if next url exists
                stop += 1
            else:
                break
        return stop

    @staticmethod
    def _check_if_url_exists_and_is_valid(url): # ConnectionError happens here
        r = requests.head(url)
        return r.status_code == 200

    def get_title(self):
        soup = Ripper(self.post_url, parser='html5lib', save_path=self.save_path, refresh=self.refresh).soup
        return soup.find_all('h2')[0].text

    def _scrap_comment_for_single_page(self, page_url):
        """Return comments and commenters on a single post page

        Returns
        --------
        OrderedDict
            Dictionary of {commenter : comments}
        """

        # User posts are contained in a table with summary='posts' attribute.
        # Each commenter name is contained inside a <tr>
        # Each comment is contained in <tr> just below the name of the commenter
        soup = Ripper(page_url, parser='html5lib', save_path=self.save_path, refresh=self.refresh).soup

        # Handle supposed anomaly by decomposing all such occurrences from the tree
        for each in soup.find_all('td', class_="l pu pd"):
            each.parent.decompose()
        rows = soup.find('table', summary='posts').find_all('tr')

        output_ordered_dict = OrderedDict()
        for i in range(0, len(rows), 2):

            topic_classes = ['bold l pu', 'bold l pu nocopy'] # topic div should be either of these classes
            for class_ in topic_classes:
                try:
                    username = rows[i].find('td', class_=class_).find('a', href=True, class_=True).text.strip()
                    break
                except AttributeError:
                    pass
                username = "Nobody" # set to nobody after exhausting all options. We cannot use finally in this case

            comment_classes = ['l w pd', 'l w pd nocopy'] # comment div should be either of these classes
            for class_ in comment_classes:
                try:
                    comment_block = rows[i+1].find('td', id=True, class_=class_).find('div', class_='narrow')
                    break
                except AttributeError:
                    pass
            parsed_block = parse_comment_block(comment_block)

            # If a username already exists (i.e. a user has already commented), append an integer to the
            # present one to differentiate them.
            if username not in output_ordered_dict:
                output_ordered_dict[username] = parsed_block
            else:
                username = "{}**{}".format(username, i)
                output_ordered_dict[username] = parsed_block
        return output_ordered_dict

    def scrap_comments_for_range_of_post_pages(self, start=0, stop=1, _all_pages=False):
        """Get contents for a range of pages from start to stop"""
        if _all_pages: # since we're starting from a zero index, we have to subtract 1 from self.max_page()
            stop = self.max_page() - 1
        while start <= stop:
            next_url = "{}/{}".format(self.post_url, start)
            next_page = self._scrap_comment_for_single_page(next_url)
            yield next_page
            start += 1

    def all_commenters(self):
        """Return list of all commenters on a post"""
        # Remember we user ** to separate a username and the number of times it is appearing on a post
        return sorted([key.split("**")[0] for each in list(self.scrap_comments_for_range_of_post_pages(stop=self.max_page())) for key, value in each.items()])

    def unique_commenters(self):
        """Return list of unique commenters on a post"""
        return sorted(set(self.all_commenters()))

    def commenters_activity_summary(self):
        """Return count of number of times a user commented on a post
        To be finished..."""
        return sort_dictionary_by_value(
            Counter(self.all_commenters())
        )

class UserCommentHistory(Nairaland):
    """
    Grab a user's comment history

    Parameters
    ------------
    user_name : str
        User's name to crawl. Default is 'seun'
    """

    def __str__(self):
        return "UserCommentHistory: {}".format(self.user_post_page)

    def __init__(self, nairaland_username, refresh=True):
        super().__init__()
        self.refresh = refresh
        BASE_URL = 'https://www.nairaland.com'
        self.save_path = os.path.join(OUTPUT_DIR, 'page_rips_user')
        if not os.path.exists(self.save_path):
            os.mkdir(self.save_path)

        p = '{}/{}'.format(BASE_URL, nairaland_username.lower())
        if self._check_if_url_exists_and_is_valid(p):
            self.user_profile_page = p
            self.user_post_page = '{}/{}/posts'.format(BASE_URL, nairaland_username.lower())
        else:
            raise NonExistentNairalandUser("This user does not exist on nairaland.")

    @staticmethod
    def _check_if_url_exists_and_is_valid(url):
        r = requests.head(url)
        return r.status_code == 200

    def user_profile(self):
        """Returns a dictionary of the user's profile"""
        pass

    def max_pages(self):
        """Return number of pages of comment for user"""
        soup = Ripper(self.user_post_page, save_path=self.save_path, refresh=True).soup
        pattern = r"\<b\>\s*(\d+)\s*\<\/b\>" # pattern to search for number of pages of comments
        try:
            return int(re.search(pattern, str(soup)).group(1))
        except AttributeError:
            raise MaximumPageNotFound("Could not find max page")

    def _scrap_comment_for_single_page(self, page_url):
        """Return comments and commenters on a single post page

        Returns
        --------
        OrderedDict
            Dictionary of {section : namedtuple}
        """

        # User comments are contained in a table with neither summary nor id attribute.
        # Then follows the rows containing the section, topic, and username and the comment itself just below it
        soup = Ripper(page_url, parser='html5lib', save_path=self.save_path, refresh=self.refresh).soup

        for each in soup.find_all('td', class_="l pu pd"):
            each.parent.decompose() # remove these trees as they are unneeded
        rows = soup.find('table', id=False, summary=False).find_all('tr')

        output_ordered_dict = OrderedDict()
        for i in range(0, len(rows), 2): # go to every second row

            topic_classes = ['bold l pu', 'bold l pu nocopy']
            for class_ in topic_classes:
                try:
                    section_topic = rows[i].find('td', class_=class_).find_all('a', href=True, class_=False)
                except AttributeError:
                    pass

            comment_classes = ['l w pd', 'l w pd nocopy']
            for class_ in comment_classes:
                try:
                    comment_block = rows[i+1].find('td', class_=class_).find('div', class_='narrow')
                except AttributeError:
                    pass

            section = section_topic[0].text.strip()
            topic = section_topic[1].text.lstrip("Re:").strip()

            parsed_block = parse_comment_block(comment_block)

            Comm = namedtuple('Comment', ['topic', 'parsed_comment'])
            Comm.topic = topic
            Comm.parsed_comment = parsed_block

            if section not in output_ordered_dict:
                output_ordered_dict[section] = Comm
            else:
                section = "{}**{}".format(section, i)
                output_ordered_dict[section] = Comm
        return output_ordered_dict

    def scrap_comments_for_range_of_user_pages(self, start=0, stop=0, _all_pages=False):
        """Get contents for a range of pages from start to stop """
        if _all_pages:
            stop = self.max_pages() - 1
        while start <= stop:
            next_url = "{}/{}".format(self.user_post_page, start)
            next_page = self._scrap_comment_for_single_page(next_url)
            yield next_page
            start += 1

class TopicCollector(Nairaland):
    """
    Collect topics from a section of nairaland

    Parameters
    -----------
    section : str
        Default section is politics

    Notes
    ------
    The methods in this class are only applicable to section urls
    """

    def __str__(self):
        return "TopicCollector: {}".format(self.post_url)

    def __init__(self, section='politics'):
        super().__init__()
        self.save_path = os.path.join(OUTPUT_DIR, 'page_rips_section')
        self.section = section
        self.post_url = 'https://www.nairaland.com/{}'.format(self.section)
        if os.path.exists(self.save_path) is False:
            os.mkdir(self.save_path)

    def max_pages(self):
        """Return number of pages in this section

        Returns
        --------
        int
            The number of pages in this section
        """
        soup = Ripper(self.post_url, parser='html5lib', save_path=self.save_path, refresh=True).soup
        number = re.search(r"\(of\s*(\d+)\s*pages\)", soup.text).group(1)
        return int(number)

    def _scrap_topics_for_a_single_page(self, page_url, refresh=True):
        """
        Yield all topics on a page

        Yields
        -------
        namedtuple
            collection of 'poster', 'title', 'url', 'number of comments'
        """
        soup = Ripper(page_url, parser='html5lib', save_path=self.save_path, refresh=True).soup
        post_table = soup.find('table', id=False, summary=False)

        for td in post_table.find_all('td', id=True):
            Post = namedtuple('Post', ['poster', 'title', 'url', 'comments', 'views', 'last_commenter', 'other_meta'])

            title_component = td.find('b').find('a', href=True)
            Post.title = title_component.text.strip()
            Post.url = 'http://www.nairaland.com' + title_component.get('href').strip()

            # there is a maximum of 7 <b> tags
            meta_component = td.find('span', class_='s').find_all('b')

            Post.poster = meta_component[0].text.strip()
            Post.comments = meta_component[1].text.strip() # count includes the post itself
            # Join all other meta as a single string
            Post.views = meta_component[2].text.strip()
            Post.last_commenter = meta_component[-1].text.strip()
            Post.other_meta = " ".join([each.text.strip() for each in meta_component[3:-1]])
            yield Post

    def scrap_topics_for_range_of_pages(self, start=0, stop=0, _all_pages=False):
        """Yield all topics between 'start' and 'end' for a section

        Parameters
        -----------
        int
            Start and end values of section

        Yields
        -------
        tuple
            same yields as for titles()
        """
        if _all_pages:
            stop = self.max_pages() - 1
        while start <= stop:
            next_url = '{}/{}'.format(self.post_url, start)
            yield self._scrap_topics_for_a_single_page(next_url)
            start += 1

def export_user_comments_to_html(username=None, max_page=5):
    """Export all of a user's comments data to a html file

    Parameters
    -----------
    str
        Username
    int
        Maximum page count for user's comments (Default is 5 pages of comments)
        loop breaks if we exceed actual count
    """

    if not username:
        raise NonExistentNairalandUser("Please provide a username.")

    destination_file = os.path.join(OUTPUT_DIR, "comments_{}_{}_pages.html".format(username.lower(), max_page))
    if os.path.exists(destination_file):
        os.remove(destination_file)
    with open(destination_file, 'a+', encoding='utf-8') as f:

        # html scaffold
        f.write("<html xmlns='http://www.w3.org/1999/xhtml'>\n")
        f.write("\t<head>\n")


        # resources
        f.write("\t\t<link rel='stylesheet' href='https://stackpath.bootstrapcdn.com/bootswatch/4.1.3/superhero/bootstrap.min.css'>\n")
        f.write("\t\t<meta name='viewport' content='width=device-width, initial-scale=1, shrink-to-fit=no'>\n")
        f.write("\t\t<script src='https://stackpath.bootstrapcdn.com/bootstrap/4.1.1/js/bootstrap.min.js' integrity='sha384-smHYKdLADwkXOn1EmN1qk/HfnUcbVRZyYmZ4qpPea6sjB/pTJ0euyQp0Mk8ck+5T' crossorigin='anonymous' async></script>\n")
        f.write("<script defer src='https://use.fontawesome.com/releases/v5.0.6/js/all.js' async></script>")
        f.write("<script src='https://code.jquery.com/jquery-3.3.1.min.js' integrity='sha256-FgpCb/KJQlLNfOu91ta32o/NMZxltwRo8QtmkMRdAu8=' crossorigin='anonymous'></script>")
        # end resources

        f.write("\t\t<title>Comment history for {} - Hack Nairaland</title>\n".format(username.lower()))
        f.write("\t</head>\n")
        f.write("\t<body style='margin-bottom:5rem;'>\n")
        # navbar
        f.write("\t\t<nav class='navbar navbar-expand-lg navbar-dark bg-primary' id='topNav'>\n")
        f.write("\t\t<a class='navbar-brand' style='font-size:36px;'>Hack Nairaland</a>\n")
        f.write("\t\t<button type='button' class='navbar-toggler my-toggler' data-toggle='collapse' data-target='.navcontent'>\n")
        f.write("\t\t<span class='sr-only'>Toggle navigation</span>\n")
        f.write("\t\t<span class='navbar-toggler-icon'></span>\n")
        f.write("\t\t</button>\n")
        f.write("\t\t<div class='collapse navbar-collapse navcontent'>\n")
        f.write("\t\t<ul class='nav navbar-nav lefthand-navigation'>\n")
        f.write("\t\t<li class='nav-item'><a class='nav-link' href='#' title='Home'>Home</a></li>\n")
        f.write("\t\t</ul>\n")
        f.write("\t\t</div>\n")
        f.write("\t\t</nav>\n")
        # end navbar
        f.write("\t\t<div class='container'>\n")
        f.write("<h1>Nairaland comment history for <a href='https://nairaland.com/{0}/posts' target='_blank'>{0}</a></h1>\n".format(username))
        # breadcrumb
        f.write("\t\t<nav aria-label='breadcrumb'>\n")
        f.write("\t\t<ol class='breadcrumb'>\n")
        f.write("\t\t<li class='breadcrumb-item'><a href='#'>Home</a></li>\n")
        f.write("\t\t<li class='breadcrumb-item'>The first {} pages</li>\n".format(max_page))
        f.write("\t\t</ol>\n")
        f.write("\t\t</nav>\n")
        # end breadcrumb

        i = 1
        for page in list(UserCommentHistory(username).scrap_comments_for_range_of_user_pages(stop=max_page)):

            f.write("\t\t\t<div id='js-scroll-target{}'>\n".format(i)) # div for targeting scroll
            f.write("\t\t\t<h2><a href='#js-scroll-target{0}' class='smooth-scroll'>Page {1}</a></h2>".format(i+1, i))
            i += 1

            for section, topic_plus_comment in page.items():
                f.write("\t\t\t<h3>Section: {}</h3>\n".format(section.split('**')[0])) # remove the ** separating section and index
                f.write('\t\t\t<h4>Subject: {}</h4>'.format(topic_plus_comment.topic))

                parsed_comment = topic_plus_comment.parsed_comment
                f.write("\t\t\t<p class='text-success'>{}</p>\n".format(parsed_comment.focus_user_comment))
                quotes = parsed_comment.quotes_ordered_dict

                for username, comment in quotes.items():
                    f.write("\t\t\t\t<h4 class='text-info'>{}</h4>\n".format(username))
                    f.write("\t\t\t\t<p class='text-primary'><em>{}</em></p>\n".format(comment))
                f.write("\t\t\t<div class='dropdown-divider' style='border:1px solid white;'></div>\n")
            f.write("\t\t\t</div>\n") # finish scroll div

        # continue up page structure
        f.write("\t\t\t<p class='float-right'><a href='#topNav' class='smooth-scroll'>Back to top</a></p>\n") # back to top
        f.write("\t\t\t<p class='float-left'>Template by <a href='https://bootswatch.com/superhero/' class='smooth-scroll'>Bootswatch</a></p>\n")
        f.write("\t\t</div>\n")
        # jquery smooth scroll
        f.write("<script>\n")
        f.write("$(document).ready(function(){\n")
        f.write("\t $('.smooth-scroll').on('click', function(event) {\n")
        f.write("\t\tif (this.hash !== '') {\n")
        f.write("\t\t\tevent.preventDefault();\n")
        f.write("\t\t\tvar hash = this.hash;\n")
        f.write("\t\t\t$('html, body').animate({\n")
        f.write("\t\t\t scrollTop: $(hash).offset().top\n")
        f.write("\t\t\t}, 800, function(){\n")
        f.write("\t\t\t window.location.hash = hash;\n")
        f.write("\t\t\t});\n")
        f.write("\t\t}\n")
        f.write("\t});\n")
        f.write("});\n")
        f.write("</script>\n")
        # end jquery smooth scroll
        f.write("\t</body>\n")
        f.write("</html>")
    os.startfile(destination_file)

def export_user_comments_to_excel(username=None, max_page=5):
    """Export a user's comments to a excel file

    Parameters
    -----------
    str
        Username
    int
        Maximum page count for user's comments (Default is 5). The loop breaks if we exceed actual count
    """
    if not username:
        raise NonExistentNairalandUser("Please provide a username.")

    work_book = OP.Workbook()
    active_sheet = work_book.active
    active_sheet.title = username

    active_sheet['A1'] = "SECTION"
    active_sheet['B1'] = "TOPIC"
    active_sheet['C1'] = 'USER_COMMENT'
    active_sheet['D1'] = "QUOTED_USER"

    row_number = 2

    for page in list(UserCommentHistory(username).scrap_comments_for_range_of_user_pages(start=0, stop=1)):
        for section, topic_plus_comment in page.items():

            active_sheet.cell(row=row_number, column=1, value=section)
            active_sheet.cell(row=row_number, column=2, value=topic_plus_comment.topic)

            parsed_comment = topic_plus_comment.parsed_comment # a namedtuple instance. Multiple cells here
            active_sheet.cell(row=row_number, column=3, value=parsed_comment.focus_user_comment)

            quotes = parsed_comment.quotes_ordered_dict

            for _username, comment in quotes.items():
                user_plus_comment = "{}: {}".format(_username, comment)
                active_sheet.cell(row=row_number, column=4, value=user_plus_comment)
                row_number += 1
            row_number += 1

    destination_file = os.path.join(OUTPUT_DIR, "{}_comments_{}_pages.xlsx".format(username, max_page))
    if os.path.exists(destination_file):
        os.remove(destination_file)

    work_book.save(destination_file)
    os.startfile(destination_file)

def export_topics_to_html(section='romance', start=0, stop=3):
    """
    Writes all topics between start and end of a section to a html file
    """

    destination_file = os.path.join(OUTPUT_DIR, "{}_page_{}_{}_pages.html".format(section, start, stop))
    if os.path.exists(destination_file):
        os.remove(destination_file)
    with open(destination_file, 'a+', encoding='utf-8') as f:

        # html scaffold
        f.write("<html xmlns='http://www.w3.org/1999/xhtml'>\n")
        f.write("\t<head>\n")

        # resources
        f.write("\t\t<link rel='stylesheet' href='https://stackpath.bootstrapcdn.com/bootswatch/4.1.3/superhero/bootstrap.min.css'>\n")
        f.write("\t\t<meta name='viewport' content='width=device-width, initial-scale=1, shrink-to-fit=no'>\n")
        f.write("\t\t<script src='https://stackpath.bootstrapcdn.com/bootstrap/4.1.1/js/bootstrap.min.js' integrity='sha384-smHYKdLADwkXOn1EmN1qk/HfnUcbVRZyYmZ4qpPea6sjB/pTJ0euyQp0Mk8ck+5T' crossorigin='anonymous' async></script>\n")
        f.write("<script defer src='https://use.fontawesome.com/releases/v5.0.6/js/all.js' async></script>")
        f.write("<script src='https://code.jquery.com/jquery-3.3.1.min.js' integrity='sha256-FgpCb/KJQlLNfOu91ta32o/NMZxltwRo8QtmkMRdAu8=' crossorigin='anonymous'></script>")
        # end resources

        f.write("\t\t<title>Topics filed under {} - Hack Nairaland</title>\n".format(section))
        f.write("\t</head>\n")
        f.write("\t<body style='margin-bottom:5rem;'>\n")
        # navbar
        f.write("\t\t<nav class='navbar navbar-expand-lg navbar-dark bg-primary' id='topNav'>\n")
        f.write("\t\t<a class='navbar-brand' style='font-size:36px;'>Hack Nairaland</a>\n")
        f.write("\t\t<button type='button' class='navbar-toggler my-toggler' data-toggle='collapse' data-target='.navcontent'>\n")
        f.write("\t\t<span class='sr-only'>Toggle navigation</span>\n")
        f.write("\t\t<span class='navbar-toggler-icon'></span>\n")
        f.write("\t\t</button>\n")
        f.write("\t\t<div class='collapse navbar-collapse navcontent'>\n")
        f.write("\t\t<ul class='nav navbar-nav lefthand-navigation'>\n")
        f.write("\t\t<li class='nav-item'><a class='nav-link' href='#' title='Home'>Home</a></li>\n")
        f.write("\t\t</ul>\n")
        f.write("\t\t</div>\n")
        f.write("\t\t</nav>\n")
        # end navbar
        f.write("\t\t<div class='container'>\n")
        f.write("<h1>Topics filed under <a href='https://nairaland.com/{0}' target='_blank'>{0}</a></h1>\n".format(section))
        # breadcrumb
        f.write("\t\t<nav aria-label='breadcrumb'>\n")
        f.write("\t\t<ol class='breadcrumb'>\n")
        f.write("\t\t<li class='breadcrumb-item'><a href='#'>Home</a></li>\n")
        f.write("\t\t<li class='breadcrumb-item'>Topics filed under{}</li>\n".format(section))
        f.write("\t\t</ol>\n")
        f.write("\t\t</nav>\n")
        # end breadcrumb

        i = 1
        topics = TopicCollector(section=section)
        for page in topics.scrap_topics_for_range_of_pages(start=start, stop=stop):

            f.write("\t\t\t<div id='js-scroll-target{}'>\n".format(i)) # div for targeting scroll
            f.write("\t\t\t<h2><a href='#js-scroll-target{0}' class='smooth-scroll'>Page {1}</a></h2>".format(i+1, i))
            i += 1

            for topic in list(page):
                f.write("\t\t\t<h3><a href='{}' target='_blank'>{}</a></h3>".format(topic.url, topic.title))
                f.write("\t\t\t<h4>Posted by <a href='https://nairaland.com/{0}/topics' target='_blank'>{0}</a></h4>".format(topic.poster))
                f.write("\t\t\t<h5>{} <i class='fas fa-comment'></i> | {} <i class='fas fa-eye'></i> | Last commenter: {} | Others: {}</h5>".format(topic.comments, topic.views, topic.last_commenter, topic.other_meta))
                f.write("\t\t\t<div class='dropdown-divider' style='border:1px solid white;'></div>\n")
            f.write("\t\t\t</div>\n") # finish scroll div

        f.write("\t\t\t<p class='float-right'><a href='#topNav' class='smooth-scroll'>Back to top</a></p>\n")
        f.write("\t\t\t<p class='float-left'>Template by <a href='https://bootswatch.com/superhero/' target='_blank'>Bootswatch</a></p>\n")
        f.write("\t\t</div>\n")
        # jquery smooth scroll
        f.write("<script>\n")
        f.write("$(document).ready(function(){\n")
        f.write("\t $('.smooth-scroll').on('click', function(event) {\n")
        f.write("\t\tif (this.hash !== '') {\n")
        f.write("\t\t\tevent.preventDefault();\n")
        f.write("\t\t\tvar hash = this.hash;\n")
        f.write("\t\t\t$('html, body').animate({\n")
        f.write("\t\t\t scrollTop: $(hash).offset().top\n")
        f.write("\t\t\t}, 800, function(){\n")
        f.write("\t\t\t window.location.hash = hash;\n")
        f.write("\t\t\t});\n")
        f.write("\t\t}\n")
        f.write("\t});\n")
        f.write("});\n")
        f.write("</script>\n")
        # end jquery smooth scroll

        # finish up page structure
        f.write("\t</body>\n")
        f.write("</html>")
    os.startfile(destination_file)

def export_topics_to_excel(section='romance', start=0, stop=3):
    """Writes all topics between start and end of a section to excel"""

    work_book = OP.Workbook()
    active_sheet = work_book.active
    active_sheet.title = section

    active_sheet['A1'] = 'POSTER'
    active_sheet['B1'] = 'TITLE'
    active_sheet['C1'] = 'LINK'
    active_sheet['D1'] = 'COMMENTS'
    active_sheet['E1'] = 'VIEWS'
    active_sheet['F1'] = 'LAST COMMENTER'
    active_sheet['G1'] = 'OTHERS'

    row_number = 2

    for page in TopicCollector(section=section).scrap_topics_for_range_of_pages(start=start, stop=stop):
        for topic in list(page):
            active_sheet.cell(row=row_number, column=1, value=topic.poster)
            active_sheet.cell(row=row_number, column=2, value=topic.title)
            active_sheet.cell(row=row_number, column=3, value=topic.url)
            active_sheet.cell(row=row_number, column=4, value=topic.comments)
            active_sheet.cell(row=row_number, column=5, value=topic.views)
            active_sheet.cell(row=row_number, column=6, value=topic.last_commenter)
            active_sheet.cell(row=row_number, column=7, value=topic.other_meta)

            row_number += 1 # advance to next row

    destination_file = os.path.join(OUTPUT_DIR, "{}_page_{}_{}_pages.xlsx".format(section, start, stop))
    if os.path.exists(destination_file):
        os.remove(destination_file)
    work_book.save(destination_file)
    os.startfile(destination_file)

def export_post_docx(post_url, start=0, stop=2, _all_pages=False):
    """Export post to .docx format"""

    i = 1
    document = Document()
    post = PostCollector(post_url)
    document.add_heading(post.get_title(), 0)
    document.add_paragraph(post_url)
    for page in list(post.scrap_comments_for_range_of_post_pages(start=0, stop=2, _all_pages=_all_pages)):
        for _username, parsed_comment in page.items():
            document.add_paragraph().add_run(_username).bold = True
            document.add_paragraph(parsed_comment.focus_user_comment)
            for commenter, comment in parsed_comment.quotes_ordered_dict.items():
                document.add_paragraph().add_run(commenter).italic = True
                document.add_paragraph().add_run(comment).italic = True
        document.add_page_break()
        i += 1

    destination_file = os.path.join(OUTPUT_DIR, "post_{}.docx".format(post.get_title()))
    if os.path.exists(destination_file):
        os.remove(destination_file)
    document.save(destination_file)
    os.startfile(destination_file)

def export_post_to_markdown(post_url, start=0, stop=2, _all_pages=False):
    """Export post to markdown format"""

    post = PostCollector(post_url)
    destination_file = os.path.join(OUTPUT_DIR, "post_{}.md".format(post.get_title()))
    if os.path.exists(destination_file):
        os.remove(destination_file)

    with open(destination_file, 'a+', encoding='utf-8') as f:
        i = 1
        f.write('# {}\n\n'.format(post.get_title()))
        f.write('[{0}]({0})\n\n'.format(post_url))
        for page in list(post.scrap_comments_for_range_of_post_pages(start=0, stop=2, _all_pages=_all_pages)):
            for _username, parsed_comment in page.items():
                f.write('**{}**\n\n'.format(_username))
                f.write('{}\n\n'.format(parsed_comment.focus_user_comment))

                for commenter, comment in parsed_comment.quotes_ordered_dict.items():
                    f.write('\t{}\n\n'.format(commenter))
                    f.write('\t{}\n\n'.format(comment))
            i += 1
    os.startfile(destination_file)

if __name__ == "__main__":
    print("\n\nSee the associated Jupyter Notebook for usage instructions.")
